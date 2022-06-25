import docx
import os
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
class Info:
    def __init__(self,List_data):
        super().__init__()            
        document = docx.Document()
        picture = document.add_picture('logo.png', width=Inches(3))
        picture.alignment = 1
        Title = document.add_heading('', level=1)
        Title.add_run('Demande de Congé').bold = True
        Title.alignment = 1
        document.add_paragraph(f"""Nom Complet : {List_data[0]}""")
        document.add_paragraph(f"""Matricule : {List_data[1]}""")
        document.add_paragraph(f"""Date : AUJOURDHUIT""")
        document.add_paragraph(f"""Responsable RH : {List_data[2]}""")
        document.add_paragraph(f"""Responsable Hiérarchique : {List_data[2]}""")
        document.add_paragraph(f"""Direction : {List_data[3]}""")
        document.add_paragraph(f"""Congé demandé : {List_data[4]}""")
        document.add_paragraph(f"""Du : {List_data[5]}""")
        document.add_paragraph(f"""Jusqu’au : {List_data[6]}""")
        document.save(f"""DemandeCongé_{List_data[0]}.docx""")
        os.system("start ayoub2.docx")